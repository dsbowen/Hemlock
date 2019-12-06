"""Download

This module performs the following download processes:
1. Create a download page.
2. Create data download files.
3. Create survey view download files.
4. Zip download files and save to Google bucket.
"""

from hemlock.routes.researcher.utils import *
from hemlock.routes.researcher.login import researcher_login_required

from datetime import timedelta
from docx import Document
from docx.shared import Inches
from io import BytesIO
from itertools import chain
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from zipfile import ZipFile
import os

# Width of survey view for imgkit
SURVEY_VIEW_IMG_WIDTH = Inches(6)

"""1. Create a download page"""
@bp.route('/download', methods=['GET','POST'])
@researcher_login_required
def download():
    """Download data"""
    return render(download_page())

@researcher_page('download')
def download_page():
    """Return the Download page
    
    The download page allows researchers to select data download files and 
    survey view download files. Data download files are selected as a 
    multi-choice question. Survey view download files are selected by 
    inputing participant IDs in a free response question. IDs are delimited 
    by commas or spaces.
    """
    download_p = Page(nav=researcher_navbar(), back=False, forward=False)
    data_q = MultiChoice(download_p, text=SELECT_FILES_TXT)
    Choice(data_q, text='Metadata', value='meta')
    Choice(data_q, text='Status Log', value='status')
    Choice(data_q, text='Dataframe', value='data')
    survey_view_q = Free(download_p, text=SURVEY_VIEW_TXT)
    Validate(survey_view_q, valid_part_ids)
    Submit(survey_view_q, store_part_ids)
    btn = Download(
        download_p, 
        text='Download', 
        callback=request.url,
        download_msg='Download Complete'
    )
    HandleForm(btn, handle_download_form)
    return download_p

def valid_part_ids(survey_view_q):
    """Check that input participant IDs are valid
    
    Participant IDs must reference participants in the database.
    """
    part_ids = parse_part_ids(survey_view_q.response)
    invalid_ids = [i for i in part_ids if Participant.query.get(i) is None]
    if invalid_ids:
        if len(invalid_ids) == 1:
            return INVALID_ID.format(invalid_ids[0])
        return INVALID_IDS.format(', '.join(invalid_ids))

def store_part_ids(survey_view_q):
    """Store participant IDs in survey view question data"""
    survey_view_q.data = parse_part_ids(survey_view_q.response)

def parse_part_ids(raw_ids):
    """Parse requested participant IDs

    Raw IDs are delimited by spaces or commas. Parsed IDs are a list of 
    integers.
    """
    if raw_ids is None:
        return []
    comma_splits = raw_ids.split(',')
    space_splits = [cs.split(' ') for cs in comma_splits]
    return [id for split in space_splits for id in split if id]

def handle_download_form(btn, response):
    """Process download file selection
    
    This method first records the file selection form response and ensures 
    it is valid. If so, it adds the appropriate file creation functions to 
    the download button.
    """
    download_p = download_page()
    db.session.add(download_p)
    btn.downloads = []
    btn.create_file_functions = []
    download_p._record_response()
    if download_p._validate():
        download_p._submit()
        data_q, survey_view_q, _ = download_p.questions
        select_data(btn, data_q.data)
        select_survey_view(btn, survey_view_q.data)
        CreateFile(btn, zip_files)
    db.session.commit()

def select_data(btn, files):
    """Add file creation functions for selected data download files"""
    if files.get('meta'):
        CreateFile(btn, create_meta)
    if files.get('status'):
        CreateFile(btn, create_status)
    if files.get('data'):
        CreateFile(btn, create_data)

def select_survey_view(btn, part_ids):
    """Add file creation function for selected survey view download files"""
    if not part_ids:
        return
    CreateFile(btn, create_views, args=part_ids)

"""2. Create data download files"""
def create_meta(btn):
    """Create metadata download file"""
    stage = 'Preparing Metadata'
    yield btn.reset(stage, 0)
    add_dataframe_to_zip(btn, DataStore.query.first().meta)
    yield btn.report(stage, 100)

def create_status(btn):
    """Create participant status log download file"""
    stage = 'Preparing Status Log'
    yield btn.reset(stage, 0)
    add_dataframe_to_zip(btn, DataStore.query.first().status_log)
    yield btn.report(stage, 100)

def create_data(btn):
    """Create dataframe download file

    Some participants will have updated data. Store data from these 
    participants before creating the dataframe.
    """
    stage = 'Storing Participant Data'
    yield btn.reset(stage, 0)
    ds = DataStore.query.first()
    db.session.add(ds)
    to_store = get_parts_to_store(ds)
    db.session.add_all(to_store)
    for i, part in enumerate(to_store):
        yield btn.report(stage, 100.0*i/len(to_store))
        ds.store_participant(part)
    add_dataframe_to_zip(btn, ds.data)
    db.session.commit()
    yield btn.report(stage, 100)

def get_parts_to_store(ds):
    """Return a list of participants to store in DataStore"""
    all_parts = Participant.query.all()
    parts_stored = ds.parts_stored
    return [
        p for p in all_parts if p.updated or p not in parts_stored
    ]

def add_dataframe_to_zip(btn, data_frame):
    """Add dataframe to zip file creation function"""
    zipfunc = btn.create_file_functions[-1]
    zipfunc.args.append(data_frame.get_download_file())

"""3. Create survey view download files"""
def create_views(btn, *part_ids):
    """Create survey views for all selected participants"""
    parts = [Participant.query.get(id) for id in part_ids]
    options = Options()
    [
        options.add_argument(arg) 
        for arg in ['--disable-gpu', '--no-sandbox', '--headless']
    ]
    options.binary_location = os.environ.get('GOOGLE_CHROME_PATH')
    chromedriver_path = os.environ.get('CHROMEDRIVER_PATH')
    if chromedriver_path:
        driver = webdriver.Chrome(chromedriver_path, chrome_options=options)
    else:
        driver = webdriver.Chrome(chrome_options=options)
    gen = chain.from_iterable([create_view(btn, p, driver) for p in parts])
    for event in gen:
        yield event
    driver.close()

def create_view(btn, part, driver):
    """Create survey view for a single participant"""
    stage = 'Creating Survey View for Participant {}'.format(part.id)
    yield btn.reset(stage, 0)
    doc = Document()
    pages = part._viewing_pages
    for i, page in enumerate(pages):
        yield btn.report(stage, 100.0*i/len(pages))
        add_page_to_doc(doc, page, driver)
    add_doc_to_zip(btn, doc, part.id)
    yield btn.report(stage, 100)

def add_page_to_doc(doc, page, driver):
    """Add survey view page to survey view doc"""
    page.convert_to_abs_paths()
    driver.get('data:text/html,'+page.html)
    body = driver.find_element_by_tag_name('body')
    page_bytes = BytesIO()
    page_bytes.write(body.screenshot_as_png)
    doc.add_picture(page_bytes, width=SURVEY_VIEW_IMG_WIDTH)
    page_bytes.close()

def add_doc_to_zip(btn, doc, part_id):
    """Add survey view doc to download zip file"""
    filename = 'Participant-{}.docx'.format(part_id)
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    zipfunc = btn.create_file_functions[-1]
    zipfunc.args.append((filename, doc_bytes))

# def create_views(btn, *part_ids):
#     """Create survey views for all selected participants"""
#     parts = [Participant.query.get(id) for id in part_ids]
#     config = imgkit.config(wkhtmltoimage=os.environ.get('WKHTMLTOIMAGE'))
#     gen = chain.from_iterable([create_view(btn, p, config) for p in parts])
#     for event in gen:
#         yield event

# def create_view(btn, part, config):
#     """Create survey view for a single participant"""
#     stage = 'Creating Survey View for Participant {}'.format(part.id)
#     yield btn.reset(stage, 0)
#     doc = Document()
#     pages = part._viewing_pages
#     for i, page in enumerate(pages):
#         yield btn.report(stage, 100.0*i/len(pages))
#         add_page_to_doc(doc, page, config)
#     add_doc_to_zip(btn, doc, part.id)
#     yield btn.report(stage, 100)

# def add_page_to_doc(doc, page, config):
#     """Add survey view page to survey view doc"""
#     page.process()
#     page_bytes = BytesIO()
#     page_bytes.write(imgkit.from_string(
#         page.html, False, 
#         css=page.external_css_paths,
#         config=config, 
#         options=OPTIONS
#     ))
#     doc.add_picture(page_bytes, width=SURVEY_VIEW_IMG_WIDTH)
#     page_bytes.close()

# def add_doc_to_zip(btn, doc, part_id):
#     """Add survey view doc to download zip file"""
#     filename = 'Participant-{}.docx'.format(part_id)
#     doc_bytes = BytesIO()
#     doc.save(doc_bytes)
#     zipfunc = btn.create_file_functions[-1]
#     zipfunc.args.append((filename, doc_bytes))

"""4. Zip download files and save to Google bucket"""
def zip_files(btn, *files):
    """Zip download files
    
    `files` is a list of (download_filename, bytes) tuples.
    """
    stage = 'Zipping Files'
    yield btn.reset(stage, 0)
    zipf_bytes = BytesIO()
    zipf = ZipFile(zipf_bytes, 'w')
    for i, (filename, io) in enumerate(files):
        yield btn.report(stage, 100.0*i/len(files))
        zipf.writestr(filename, io.getvalue())
        io.close()
    zipf.close()
    blob = current_app.gcp_bucket.blob(btn.model_id+'.zip')
    blob.upload_from_string(zipf_bytes.getvalue())
    url = blob.generate_signed_url(expiration=timedelta(hours=1))
    btn.downloads = [(url, 'download.zip')]
    zipf_bytes.close()
    yield btn.report(stage, 100)